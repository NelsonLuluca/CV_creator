#import modules
from docx import Document
from docx.shared import Inches
#import pyttsx3


document = Document()

#add photo section
document.add_picture('LulucaDev.JPG', width=Inches(2.0))

#name, phone number and email section
name = input("What's your name?\n")
phone_number = (input("What's your phone number\n"))
email = input("What's your e-mail\n")

#concatenate everything section
document.add_paragraph(name + '|' +phone_number + '|' + email )

#about me section
document.add_heading('About me')
about_me = input('Tell me about yourself?\n ')

document.add_paragraph(about_me)

#Work experience
document.add_heading('Work experience')
p = document.add_paragraph()

company = input('Enter Company \n')
from_date = input('From Date \n')
to_date = input('To Date \n')

p.add_run(company + ' ').bolt=True 
p.add_run(from_date + '-'+ to_date + ' \n').italic=True

experience_details = input(f'Describe your experience at {company} \n')
p.add_run(experience_details)

#more experiences loop
while True:
    has_more_experiences = input("Do you have more experiences? YES or NO \n")
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter Company \n')
        from_date = input('From Date \n')
        to_date = input('To Date \n')

        p.add_run(company + ' ').bolt=True 
        p.add_run(from_date + '-'+ to_date + ' \n').italic=True

        experience_details = input(f'Describe your experience at {company} \n')
        p.add_run(experience_details)
    else:
        break
    
    #skills
    document.add_heading('Skills')
    skill = input('Enter your skill \n')
    p = document.add_paragraph(skill)
    p.style = 'List Bullet'
    
    #skill loop
    while True:
        has_more_skills = input('Do you have more skill? \n')
        if has_more_skills.lower() == 'yes':
             skill = input('Enter your skill \n')
             p = document.add_paragraph(skill)
             p.style = 'List Bullet'
        else:
            break
    
#footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using NelsonLuluca CVMAKER"
            


document.save("cv.docx")

